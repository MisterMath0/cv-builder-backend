# app/services/export_service.py
from fastapi import HTTPException
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
import logging
from typing import Dict, Any, List
from datetime import datetime

logger = logging.getLogger(__name__)

class ExportService:
    async def to_pdf(self, html_content: str) -> bytes:
        """Convert HTML to PDF"""
        try:
            logger.debug("Converting HTML to PDF")
            pdf = HTML(string=html_content).write_pdf()
            logger.debug("PDF generated successfully")
            return pdf
        except Exception as e:
            logger.error(f"PDF generation failed: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to generate PDF: {str(e)}"
            )

    def parse_date(self, date_string: str) -> str:
        """Parse date string to formatted date"""
        if not date_string:
            return ""
        try:
            date_string = date_string.split('T')[0]
            date = datetime.strptime(date_string, '%Y-%m-%d')
            return date.strftime('%B %Y')
        except Exception:
            return date_string

    async def to_docx(self, cv_data: Dict[str, Any]) -> bytes:
        """Convert CV data to DOCX"""
        try:
            logger.debug("Generating DOCX document")
            doc = Document()
            
            # Document styling
            self._setup_document_style(doc)
            
            # Process sections in order
            sections = cv_data.get('sections', [])
            for section in sorted(sections, key=lambda x: x.get('order', 0)):
                section_type = section.get('type')
                content = section.get('content')
                title = section.get('title')

                if section_type == 'contact':
                    self._add_contact_section(doc, content)
                elif section_type == 'text':
                    self._add_section(doc, title, content)
                elif section_type == 'experience':
                    self._add_experience_section(doc, content, title)
                elif section_type == 'education':
                    self._add_education_section(doc, content, title)
                elif section_type == 'skills':
                    self._add_section(doc, title, content)
                elif section_type == 'languages':
                    self._add_languages_section(doc, content, title)
                elif section_type == 'hobbies':
                    self._add_section(doc, title, content)

            # Save to bytes
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            return docx_buffer.read()
            
        except Exception as e:
            logger.error(f"DOCX generation failed: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to generate DOCX: {str(e)}"
            )

    def _setup_document_style(self, doc: Document):
        """Setup document styling"""
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        for i in range(1, 4):
            style = doc.styles[f'Heading {i}']
            style.font.name = 'Calibri'
            style.font.size = Pt(16 - i)
            style.font.bold = True

    def _add_contact_section(self, doc: Document, contact_data: Dict):
        """Add contact information section"""
        if not contact_data:
            return

        name = contact_data.get('name', '')
        email = contact_data.get('email', '')
        phone = contact_data.get('phone', '')
        location = contact_data.get('location', '')
        
        if name:
            heading = doc.add_heading(name, 0)
            heading.alignment = 1

        contact_para = doc.add_paragraph()
        contact_para.alignment = 1
        if email:
            contact_para.add_run(email)
        if phone:
            contact_para.add_run(f" | {phone}" if email else phone)
        if location:
            contact_para.add_run(f"\n{location}")
        
        doc.add_paragraph()

    def _add_section(self, doc: Document, title: str, content: str):
        """Add a generic section"""
        if not content:
            return
            
        doc.add_heading(title, 1)
        doc.add_paragraph(content)
        doc.add_paragraph()

    def _add_experience_section(self, doc: Document, experiences: List[Dict], title: str):
        """Add experience section"""
        if not experiences:
            return
            
        doc.add_heading(title, 1)
        
        for exp in experiences:
            p = doc.add_paragraph()
            if exp.get('position'):
                p.add_run(f"{exp['position']}").bold = True
            if exp.get('company'):
                p.add_run(f" at {exp['company']}\n").italic = True
            
            start_date = self.parse_date(exp.get('startDate'))
            if start_date:
                date_text = f"{start_date} - "
                date_text += "Present" if exp.get('current') else self.parse_date(exp.get('endDate', ''))
                p.add_run(date_text + '\n').italic = True
            
            if description := exp.get('description'):
                doc.add_paragraph(description)
            
            doc.add_paragraph()

    def _add_education_section(self, doc: Document, education: List[Dict], title: str):
        """Add education section"""
        if not education:
            return
            
        doc.add_heading(title, 1)
        
        for edu in education:
            p = doc.add_paragraph()
            if edu.get('degree'):
                p.add_run(f"{edu['degree']}\n").bold = True
            if edu.get('institution'):
                p.add_run(f"{edu['institution']}\n").italic = True
            
            start_date = self.parse_date(edu.get('startDate'))
            if start_date:
                date_text = f"{start_date} - "
                date_text += "Present" if edu.get('current') else self.parse_date(edu.get('endDate', ''))
                p.add_run(date_text + '\n').italic = True
            
            if description := edu.get('description'):
                doc.add_paragraph(description)
            
            doc.add_paragraph()

    def _add_languages_section(self, doc: Document, languages: List[Dict], title: str):
        """Add languages section"""
        if not languages:
            return
            
        doc.add_heading(title, 1)
        
        for lang in languages:
            if lang.get('name') and lang.get('level'):
                doc.add_paragraph(
                    f"{lang['name']} - {lang['level']}", 
                    style='List Bullet'
                )
        
        doc.add_paragraph()