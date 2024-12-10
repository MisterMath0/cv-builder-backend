from fastapi import HTTPException
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import jinja2
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)

class CoverLetterExportService:
    def __init__(self):
        try:
            base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            template_dir = os.path.join(base_dir, 'templates', 'cover_letter')
            
            self.template_loader = jinja2.FileSystemLoader(template_dir)
            self.template_env = jinja2.Environment(
                loader=self.template_loader,
                autoescape=True
            )
        except Exception as e:
            logger.error(f"Template environment initialization failed: {str(e)}")
            raise

    async def to_pdf(
        self, 
        content: str,
        company_name: Optional[str] = None,
        job_title: Optional[str] = None,
        author: Optional[str] = None,
        template_name: str = "basic.html"
    ) -> bytes:
        """Convert cover letter to PDF using template"""
        try:
            template = self.template_env.get_template(template_name)
            
            html_content = template.render(
                content=content,
                company_name=company_name,
                job_title=job_title,
                author=author,
                date=datetime.now().strftime("%B %d, %Y")
            )
            
            pdf = HTML(string=html_content).write_pdf()
            return pdf
        except Exception as e:
            logger.error(f"PDF generation failed: {str(e)}")
            raise HTTPException(status_code=400, detail=str(e))

    async def to_docx(
        self, 
        content: str,
        company_name: Optional[str] = None,
        job_title: Optional[str] = None,
        author: Optional[str] = None
    ) -> bytes:
        """Convert cover letter to DOCX using template-based content"""
        try:
            # First generate HTML using template for consistent formatting
            template = self.template_env.get_template("basic.html")
            html_content = template.render(
                content=content,
                company_name=company_name,
                job_title=job_title,
                author=author,
                date=datetime.now().strftime("%B %d, %Y")
            )
            
            # Create Word document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Add date
            date_paragraph = doc.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_paragraph.add_run(datetime.now().strftime("%B %d, %Y"))
            date_run.font.size = Pt(11)

            # Add company info
            if company_name:
                company_para = doc.add_paragraph()
                company_para.add_run(company_name + "\n").bold = True
                if job_title:
                    company_para.add_run(f"Re: {job_title}\n")
                doc.add_paragraph()

            # Add content
            content_para = doc.add_paragraph()
            content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = content_para.add_run(content)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

            # Add author if provided
            if author:
                doc.add_paragraph()
                author_para = doc.add_paragraph()
                author_para.add_run(author)

            # Convert to bytes
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            return docx_buffer.read()

        except Exception as e:
            logger.error(f"DOCX generation failed: {str(e)}")
            raise HTTPException(status_code=400, detail=str(e))