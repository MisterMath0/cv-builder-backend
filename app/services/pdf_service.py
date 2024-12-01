# app/services/pdf_service.py
from weasyprint import HTML
import logging
from fastapi import HTTPException
from docx import Document
from typing import Dict, Any

logger = logging.getLogger(__name__)

class PDFService:
    async def generate_pdf(self, html_content: str) -> bytes:
        """Convert HTML preview to PDF"""
        try:
            pdf = HTML(string=html_content).write_pdf()
            return pdf
        except Exception as e:
            logger.error(f"Failed to generate PDF: {str(e)}")
            raise HTTPException(status_code=400, detail=f"PDF generation failed: {str(e)}")

    async def generate_docx(self, cv_data: Dict[Any, Any]) -> bytes:
        """Generate Word document from CV data"""
        try:
            doc = Document()
            
            # Add sections systematically
            sections_order = ['contact', 'profile', 'experience', 'education', 'skills', 'languages', 'hobbies']
            
            for section_type in sections_order:
                if section_data := cv_data.get(section_type):
                    self._add_section(doc, section_type, section_data)

            doc_bytes = bytes()
            doc.save(doc_bytes)
            return doc_bytes
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {str(e)}")
            raise HTTPException(status_code=400, detail=f"DOCX generation failed: {str(e)}")

    def _add_section(self, doc: Document, section_type: str, data: Any):
        """Add a section to the document based on its type"""
        section_handlers = {
            'contact': self._add_contact_section,
            'profile': lambda d, data: self._add_text_section(d, "Profile", data),
            'experience': self._add_experience_section,
            'education': self._add_education_section,
            'skills': lambda d, data: self._add_text_section(d, "Skills", data),
            'languages': self._add_languages_section,
            'hobbies': lambda d, data: self._add_text_section(d, "Hobbies & Interests", data)
        }
        
        if handler := section_handlers.get(section_type):
            handler(doc, data)

    def _add_contact_section(self, doc: Document, data: dict):
        doc.add_heading('Contact Information', 1)
        p = doc.add_paragraph()
        p.add_run(f"{data.get('name', '')}\n").bold = True
        p.add_run(f"Email: {data.get('email', '')}\n")
        p.add_run(f"Phone: {data.get('phone', '')}\n")
        p.add_run(f"Location: {data.get('location', '')}")

    def _add_text_section(self, doc: Document, title: str, content: str):
        doc.add_heading(title, 1)
        doc.add_paragraph(content)

    def _add_experience_section(self, doc: Document, experiences: list):
        doc.add_heading('Professional Experience', 1)
        for exp in experiences:
            p = doc.add_paragraph()
            p.add_run(f"{exp['position']} at {exp['company']}\n").bold = True
            p.add_run(f"{exp['startDate']} - {exp['endDate'] if not exp['current'] else 'Present'}\n").italic = True
            doc.add_paragraph(exp['description'])

    def _add_education_section(self, doc: Document, education: list):
        doc.add_heading('Education', 1)
        for edu in education:
            p = doc.add_paragraph()
            p.add_run(f"{edu['degree']} - {edu['institution']}\n").bold = True
            p.add_run(f"{edu['startDate']} - {edu['endDate'] if not edu['current'] else 'Present'}\n").italic = True
            doc.add_paragraph(edu['description'])

    def _add_languages_section(self, doc: Document, languages: list):
        doc.add_heading('Languages', 1)
        for lang in languages:
            doc.add_paragraph(f"{lang['name']} - {lang['level']}", style='List Bullet')